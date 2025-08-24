import json
import os
from typing import List, Dict, Optional, Any
from datetime import datetime
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from io import BytesIO
import base64
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
import plotly.graph_objects as go
import plotly.express as px
from plotly.subplots import make_subplots
import statistics

class ReportGenerator:
    """Comprehensive report generation system with multiple formats and visualizations"""
    
    def __init__(self):
        self.report_templates_dir = "grading_reports/templates"
        self.reports_dir = "grading_reports"
        self._ensure_directories()
        
    def _ensure_directories(self):
        """Ensure required directories exist"""
        os.makedirs(self.reports_dir, exist_ok=True)
        os.makedirs(self.report_templates_dir, exist_ok=True)
        os.makedirs(f"{self.reports_dir}/individual", exist_ok=True)
        os.makedirs(f"{self.reports_dir}/class", exist_ok=True)
        os.makedirs(f"{self.reports_dir}/json", exist_ok=True)
        os.makedirs(f"{self.reports_dir}/charts", exist_ok=True)
    
    def generate_individual_report_docx(self, grading_result, student_info: Dict, 
                                      test_info: Dict) -> str:
        """Generate individual student report in Word format"""
        try:
            doc = Document()
            
            # Add title page
            self._add_title_page(doc, student_info, test_info, grading_result)
            
            # Add executive summary
            self._add_executive_summary(doc, grading_result)
            
            # Add detailed question analysis
            self._add_question_analysis(doc, grading_result)
            
            # Add performance charts
            self._add_performance_charts(doc, grading_result)
            
            # Add recommendations
            self._add_recommendations(doc, grading_result)
            
            # Add rubric compliance
            self._add_rubric_compliance(doc, grading_result)
            
            # Save document
            filename = f"{self.reports_dir}/individual/{student_info['name']}_{test_info['title']}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            doc.save(filename)
            
            return filename
            
        except Exception as e:
            print(f"Error generating Word report: {e}")
            raise
    
    def _add_title_page(self, doc: Document, student_info: Dict, test_info: Dict, 
                       grading_result) -> None:
        """Add professional title page"""
        # Title
        title = doc.add_heading('Student Performance Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Student and test information
        doc.add_paragraph()
        info_table = doc.add_table(rows=4, cols=2)
        info_table.style = 'Table Grid'
        
        # Student info
        info_table.cell(0, 0).text = "Student Name:"
        info_table.cell(0, 1).text = student_info.get('name', 'N/A')
        info_table.cell(1, 0).text = "Student ID:"
        info_table.cell(1, 1).text = student_info.get('student_id', 'N/A')
        
        # Test info
        info_table.cell(2, 0).text = "Test:"
        info_table.cell(2, 1).text = test_info.get('title', 'N/A')
        info_table.cell(3, 0).text = "Date:"
        info_table.cell(3, 1).text = datetime.now().strftime('%B %d, %Y')
        
        # Overall score
        doc.add_paragraph()
        score_para = doc.add_paragraph()
        score_para.add_run(f"Overall Score: {grading_result.percentage:.1f}% ").bold = True
        score_para.add_run(f"({grading_result.total_score:.1f}/{grading_result.max_possible_score:.1f})")
        
        # Grading confidence
        confidence_para = doc.add_paragraph()
        confidence_para.add_run(f"Grading Confidence: {grading_result.grading_confidence:.1%}")
        
        doc.add_page_break()
    
    def _add_executive_summary(self, doc: Document, grading_result) -> None:
        """Add executive summary section"""
        doc.add_heading('Executive Summary', level=1)
        
        # Performance overview
        doc.add_heading('Performance Overview', level=2)
        doc.add_paragraph(f"Overall Performance: {grading_result.percentage:.1f}%")
        doc.add_paragraph(f"Total Questions: {len(grading_result.question_scores)}")
        doc.add_paragraph(f"Grading Confidence: {grading_result.grading_confidence:.1%}")
        
        # Key strengths
        if grading_result.strengths:
            doc.add_heading('Key Strengths', level=2)
            for strength in grading_result.strengths:
                doc.add_paragraph(f"• {strength}", style='List Bullet')
        
        # Areas for improvement
        if grading_result.areas_for_improvement:
            doc.add_heading('Areas for Improvement', level=2)
            for area in grading_result.areas_for_improvement:
                doc.add_paragraph(f"• {area}", style='List Bullet')
        
        doc.add_paragraph()
    
    def _add_question_analysis(self, doc: Document, grading_result) -> None:
        """Add detailed question-by-question analysis"""
        doc.add_heading('Detailed Question Analysis', level=1)
        
        for qs in grading_result.question_scores:
            # Question header
            doc.add_heading(f'Question {qs.question_number}', level=2)
            
            # Question score table
            score_table = doc.add_table(rows=1, cols=4)
            score_table.style = 'Table Grid'
            score_table.cell(0, 0).text = "Awarded Marks"
            score_table.cell(0, 1).text = "Total Marks"
            score_table.cell(0, 2).text = "Percentage"
            score_table.cell(0, 3).text = "Confidence"
            
            score_row = score_table.add_row()
            score_row.cells[0].text = f"{qs.awarded_marks:.1f}"
            score_row.cells[1].text = f"{qs.total_marks:.1f}"
            score_row.cells[2].text = f"{qs.percentage:.1f}%"
            score_row.cells[3].text = f"{qs.confidence:.1%}"
            
            # Component scores
            doc.add_heading('Component Scores', level=3)
            comp_table = doc.add_table(rows=1, cols=3)
            comp_table.style = 'Table Grid'
            comp_table.cell(0, 0).text = "Mathematical Reasoning"
            comp_table.cell(0, 1).text = "Conceptual Understanding"
            comp_table.cell(0, 2).text = "Presentation"
            
            comp_row = comp_table.add_row()
            comp_row.cells[0].text = f"{qs.mathematical_reasoning_score:.1%}"
            comp_row.cells[1].text = f"{qs.conceptual_understanding_score:.1%}"
            comp_row.cells[2].text = f"{qs.presentation_score:.1%}"
            
            # Step-by-step evaluation
            if qs.step_evaluations:
                doc.add_heading('Step-by-Step Evaluation', level=3)
                step_table = doc.add_table(rows=1, cols=5)
                step_table.style = 'Table Grid'
                step_table.cell(0, 0).text = "Step"
                step_table.cell(0, 1).text = "Description"
                step_table.cell(0, 2).text = "Correct"
                step_table.cell(0, 3).text = "Partial Credit"
                step_table.cell(0, 4).text = "Feedback"
                
                for step in qs.step_evaluations:
                    step_row = step_table.add_row()
                    step_row.cells[0].text = str(step.step_number)
                    step_row.cells[1].text = step.step_description[:50] + "..." if len(step.step_description) > 50 else step.step_description
                    step_row.cells[2].text = "✓" if step.is_correct else "✗"
                    step_row.cells[3].text = f"{step.partial_credit:.1%}"
                    step_row.cells[4].text = step.feedback[:50] + "..." if len(step.feedback) > 50 else step.feedback
            
            # Feedback
            doc.add_heading('Feedback', level=3)
            doc.add_paragraph(qs.overall_feedback)
            
            doc.add_paragraph()
    
    def _add_performance_charts(self, doc: Document, grading_result) -> None:
        """Add performance visualization charts"""
        doc.add_heading('Performance Visualizations', level=1)
        
        # Generate charts
        charts = self._generate_individual_charts(grading_result)
        
        for chart_name, chart_data in charts.items():
            doc.add_heading(chart_name, level=2)
            
            # Add chart image
            if chart_data:
                doc.add_picture(chart_data, width=Inches(6))
                doc.add_paragraph()
    
    def _add_recommendations(self, doc: Document, grading_result) -> None:
        """Add recommendations section"""
        doc.add_heading('Recommendations', level=1)
        
        # Overall recommendations
        doc.add_heading('Overall Recommendations', level=2)
        doc.add_paragraph(grading_result.overall_feedback)
        
        # Question-specific recommendations
        doc.add_heading('Question-Specific Recommendations', level=2)
        for qs in grading_result.question_scores:
            if qs.suggestions:
                doc.add_heading(f'Question {qs.question_number}', level=3)
                for suggestion in qs.suggestions:
                    doc.add_paragraph(f"• {suggestion}", style='List Bullet')
        
        doc.add_paragraph()
    
    def _add_rubric_compliance(self, doc: Document, grading_result) -> None:
        """Add rubric compliance analysis"""
        doc.add_heading('Rubric Compliance Analysis', level=1)
        
        if grading_result.rubric_compliance:
            comp_table = doc.add_table(rows=1, cols=2)
            comp_table.style = 'Table Grid'
            comp_table.cell(0, 0).text = "Criteria"
            comp_table.cell(0, 1).text = "Compliance Score"
            
            for criteria, score in grading_result.rubric_compliance.items():
                row = comp_table.add_row()
                row.cells[0].text = criteria.replace('_', ' ').title()
                row.cells[1].text = f"{score:.1%}"
        
        doc.add_paragraph()
    
    def _generate_individual_charts(self, grading_result) -> Dict[str, BytesIO]:
        """Generate charts for individual report"""
        charts = {}
        
        try:
            # Question scores chart
            fig, ax = plt.subplots(figsize=(10, 6))
            questions = [f"Q{qs.question_number}" for qs in grading_result.question_scores]
            scores = [qs.percentage for qs in grading_result.question_scores]
            
            bars = ax.bar(questions, scores, color=['green' if s >= 80 else 'orange' if s >= 60 else 'red' for s in scores])
            ax.set_ylabel('Percentage Score')
            ax.set_title('Question-by-Question Performance')
            ax.set_ylim(0, 100)
            
            # Add value labels on bars
            for bar, score in zip(bars, scores):
                height = bar.get_height()
                ax.text(bar.get_x() + bar.get_width()/2., height + 1,
                       f'{score:.1f}%', ha='center', va='bottom')
            
            plt.tight_layout()
            
            # Save to BytesIO
            chart_buffer = BytesIO()
            plt.savefig(chart_buffer, format='png', dpi=300, bbox_inches='tight')
            chart_buffer.seek(0)
            charts['Question Performance Chart'] = chart_buffer
            plt.close()
            
            # Component scores radar chart
            if grading_result.question_scores:
                fig, ax = plt.subplots(figsize=(8, 8), subplot_kw=dict(projection='polar'))
                
                categories = ['Mathematical\nReasoning', 'Conceptual\nUnderstanding', 'Presentation']
                avg_scores = [
                    statistics.mean([qs.mathematical_reasoning_score for qs in grading_result.question_scores]),
                    statistics.mean([qs.conceptual_understanding_score for qs in grading_result.question_scores]),
                    statistics.mean([qs.presentation_score for qs in grading_result.question_scores])
                ]
                
                angles = [n / float(len(categories)) * 2 * 3.14159 for n in range(len(categories))]
                angles += angles[:1]  # Complete the circle
                avg_scores += avg_scores[:1]
                
                ax.plot(angles, avg_scores, 'o-', linewidth=2, label='Student Performance')
                ax.fill(angles, avg_scores, alpha=0.25)
                ax.set_xticks(angles[:-1])
                ax.set_xticklabels(categories)
                ax.set_ylim(0, 1)
                ax.set_title('Component Skills Analysis', pad=20)
                
                plt.tight_layout()
                
                # Save to BytesIO
                radar_buffer = BytesIO()
                plt.savefig(radar_buffer, format='png', dpi=300, bbox_inches='tight')
                radar_buffer.seek(0)
                charts['Component Skills Radar Chart'] = radar_buffer
                plt.close()
            
        except Exception as e:
            print(f"Error generating charts: {e}")
        
        return charts
    
    def generate_class_report_docx(self, submissions_data: List[Dict], test_info: Dict) -> str:
        """Generate class-wide performance report"""
        try:
            doc = Document()
            
            # Title page
            doc.add_heading(f'Class Performance Report - {test_info["title"]}', 0)
            doc.add_paragraph(f"Generated on: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
            doc.add_paragraph(f"Total Students: {len(submissions_data)}")
            
            # Class statistics
            self._add_class_statistics(doc, submissions_data)
            
            # Performance distribution
            self._add_performance_distribution(doc, submissions_data)
            
            # Question analysis
            self._add_class_question_analysis(doc, submissions_data)
            
            # Student rankings
            self._add_student_rankings(doc, submissions_data)
            
            # Save document
            filename = f"{self.reports_dir}/class/{test_info['title']}_ClassReport_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            doc.save(filename)
            
            return filename
            
        except Exception as e:
            print(f"Error generating class report: {e}")
            raise
    
    def _add_class_statistics(self, doc: Document, submissions_data: List[Dict]) -> None:
        """Add class statistics section"""
        doc.add_heading('Class Statistics', level=1)
        
        # Calculate statistics
        scores = [sub['percentage'] for sub in submissions_data if sub.get('percentage')]
        if scores:
            stats_table = doc.add_table(rows=6, cols=2)
            stats_table.style = 'Table Grid'
            
            stats_data = [
                ("Mean Score", f"{statistics.mean(scores):.1f}%"),
                ("Median Score", f"{statistics.median(scores):.1f}%"),
                ("Standard Deviation", f"{statistics.stdev(scores):.1f}%"),
                ("Highest Score", f"{max(scores):.1f}%"),
                ("Lowest Score", f"{min(scores):.1f}%"),
                ("Pass Rate (≥60%)", f"{sum(1 for s in scores if s >= 60) / len(scores) * 100:.1f}%")
            ]
            
            for i, (stat, value) in enumerate(stats_data):
                stats_table.cell(i, 0).text = stat
                stats_table.cell(i, 1).text = value
        
        doc.add_paragraph()
    
    def _add_performance_distribution(self, doc: Document, submissions_data: List[Dict]) -> None:
        """Add performance distribution analysis"""
        doc.add_heading('Performance Distribution', level=1)
        
        scores = [sub['percentage'] for sub in submissions_data if sub.get('percentage')]
        if scores:
            # Create distribution chart
            fig, ax = plt.subplots(figsize=(10, 6))
            ax.hist(scores, bins=10, edgecolor='black', alpha=0.7)
            ax.set_xlabel('Score (%)')
            ax.set_ylabel('Number of Students')
            ax.set_title('Score Distribution')
            ax.axvline(statistics.mean(scores), color='red', linestyle='--', label=f'Mean: {statistics.mean(scores):.1f}%')
            ax.legend()
            
            plt.tight_layout()
            
            # Save and add to document
            chart_buffer = BytesIO()
            plt.savefig(chart_buffer, format='png', dpi=300, bbox_inches='tight')
            chart_buffer.seek(0)
            doc.add_picture(chart_buffer, width=Inches(6))
            plt.close()
        
        doc.add_paragraph()
    
    def _add_class_question_analysis(self, doc: Document, submissions_data: List[Dict]) -> None:
        """Add class-wide question analysis"""
        doc.add_heading('Question Analysis', level=1)
        
        # This would require detailed question data from each submission
        # For now, add a placeholder
        doc.add_paragraph("Detailed question analysis requires individual question scores from all submissions.")
        doc.add_paragraph()
    
    def _add_student_rankings(self, doc: Document, submissions_data: List[Dict]) -> None:
        """Add student rankings table"""
        doc.add_heading('Student Rankings', level=1)
        
        # Sort by percentage
        sorted_submissions = sorted(submissions_data, key=lambda x: x.get('percentage', 0), reverse=True)
        
        if sorted_submissions:
            rank_table = doc.add_table(rows=1, cols=4)
            rank_table.style = 'Table Grid'
            rank_table.cell(0, 0).text = "Rank"
            rank_table.cell(0, 1).text = "Student Name"
            rank_table.cell(0, 2).text = "Score (%)"
            rank_table.cell(0, 3).text = "Total Marks"
            
            for i, submission in enumerate(sorted_submissions, 1):
                row = rank_table.add_row()
                row.cells[0].text = str(i)
                row.cells[1].text = submission.get('student_name', 'N/A')
                row.cells[2].text = f"{submission.get('percentage', 0):.1f}%"
                row.cells[3].text = f"{submission.get('total_score', 0):.1f}/{submission.get('max_possible_score', 0):.1f}"
        
        doc.add_paragraph()
    
    def export_json_data(self, grading_result, student_info: Dict, test_info: Dict) -> str:
        """Export grading data as JSON for external systems"""
        try:
            export_data = {
                "metadata": {
                    "export_date": datetime.now().isoformat(),
                    "student_info": student_info,
                    "test_info": test_info,
                    "grading_metadata": {
                        "grading_time": grading_result.grading_time.isoformat(),
                        "grading_confidence": grading_result.grading_confidence,
                        "total_questions": len(grading_result.question_scores)
                    }
                },
                "overall_performance": {
                    "total_score": grading_result.total_score,
                    "max_possible_score": grading_result.max_possible_score,
                    "percentage": grading_result.percentage,
                    "overall_feedback": grading_result.overall_feedback,
                    "strengths": grading_result.strengths,
                    "areas_for_improvement": grading_result.areas_for_improvement
                },
                "question_scores": [
                    {
                        "question_number": qs.question_number,
                        "total_marks": qs.total_marks,
                        "awarded_marks": qs.awarded_marks,
                        "percentage": qs.percentage,
                        "confidence": qs.confidence,
                        "component_scores": {
                            "mathematical_reasoning": qs.mathematical_reasoning_score,
                            "conceptual_understanding": qs.conceptual_understanding_score,
                            "presentation": qs.presentation_score
                        },
                        "step_evaluations": [
                            {
                                "step_number": step.step_number,
                                "description": step.step_description,
                                "is_correct": step.is_correct,
                                "partial_credit": step.partial_credit,
                                "feedback": step.feedback,
                                "reasoning": step.reasoning,
                                "confidence": step.confidence
                            }
                            for step in qs.step_evaluations
                        ],
                        "feedback": {
                            "overall": qs.overall_feedback,
                            "strengths": qs.strengths,
                            "weaknesses": qs.weaknesses,
                            "suggestions": qs.suggestions
                        }
                    }
                    for qs in grading_result.question_scores
                ],
                "rubric_compliance": grading_result.rubric_compliance,
                "performance_analysis": grading_result.performance_analysis
            }
            
            filename = f"{self.reports_dir}/json/{student_info['name']}_{test_info['title']}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
            
            with open(filename, 'w', encoding='utf-8') as f:
                json.dump(export_data, f, indent=2, ensure_ascii=False)
            
            return filename
            
        except Exception as e:
            print(f"Error exporting JSON data: {e}")
            raise
    
    def generate_interactive_charts(self, grading_result, student_info: Dict) -> Dict[str, str]:
        """Generate interactive Plotly charts"""
        charts = {}
        
        try:
            # Question performance bar chart
            questions = [f"Q{qs.question_number}" for qs in grading_result.question_scores]
            scores = [qs.percentage for qs in grading_result.question_scores]
            
            fig = go.Figure(data=[
                go.Bar(
                    x=questions,
                    y=scores,
                    marker_color=['green' if s >= 80 else 'orange' if s >= 60 else 'red' for s in scores],
                    text=[f'{s:.1f}%' for s in scores],
                    textposition='auto'
                )
            ])
            
            fig.update_layout(
                title=f'Question Performance - {student_info["name"]}',
                xaxis_title='Questions',
                yaxis_title='Score (%)',
                yaxis=dict(range=[0, 100])
            )
            
            charts['question_performance'] = fig.to_html(include_plotlyjs=False, full_html=False)
            
            # Component skills radar chart
            if grading_result.question_scores:
                categories = ['Mathematical Reasoning', 'Conceptual Understanding', 'Presentation']
                avg_scores = [
                    statistics.mean([qs.mathematical_reasoning_score for qs in grading_result.question_scores]),
                    statistics.mean([qs.conceptual_understanding_score for qs in grading_result.question_scores]),
                    statistics.mean([qs.presentation_score for qs in grading_result.question_scores])
                ]
                
                fig = go.Figure()
                fig.add_trace(go.Scatterpolar(
                    r=avg_scores,
                    theta=categories,
                    fill='toself',
                    name='Student Performance'
                ))
                
                fig.update_layout(
                    polar=dict(
                        radialaxis=dict(
                            visible=True,
                            range=[0, 1]
                        )),
                    showlegend=False,
                    title=f'Component Skills Analysis - {student_info["name"]}'
                )
                
                charts['component_skills'] = fig.to_html(include_plotlyjs=False, full_html=False)
            
            # Step evaluation timeline
            all_steps = []
            for qs in grading_result.question_scores:
                for step in qs.step_evaluations:
                    all_steps.append({
                        'question': qs.question_number,
                        'step': step.step_number,
                        'credit': step.partial_credit,
                        'correct': step.is_correct
                    })
            
            if all_steps:
                df = pd.DataFrame(all_steps)
                fig = px.scatter(
                    df, x='question', y='step', 
                    color='credit', size='credit',
                    title=f'Step-by-Step Performance - {student_info["name"]}',
                    labels={'question': 'Question Number', 'step': 'Step Number', 'credit': 'Partial Credit'}
                )
                
                charts['step_performance'] = fig.to_html(include_plotlyjs=False, full_html=False)
            
        except Exception as e:
            print(f"Error generating interactive charts: {e}")
        
        return charts

# Global report generator instance
report_generator = ReportGenerator()
